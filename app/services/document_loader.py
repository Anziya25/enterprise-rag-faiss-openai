"""
Document Loader

Responsible for reading different document formats
and converting them into LangChain Document objects.
"""

from pathlib import Path
from typing import List

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain_core.documents import Document


class DocumentLoader:
    """
    Loads TXT, PDF and DOCX files and converts them
    into LangChain Document objects.
    """

    def load_txt(self, file_path: str) -> List[Document]:
        """
        Load a TXT file.
        """

        path = Path(file_path)

        with open(path, "r", encoding="utf-8") as file:
            text = file.read()

        return [
            Document(
                page_content=text,
                metadata={
                    "source": path.name,
                    "file_type": ".txt"
                }
            )
        ]

    def load_pdf(self, file_path: str) -> List[Document]:
        """
        Load a PDF file.

        Returns one LangChain Document per page.
        """

        path = Path(file_path)

        pdf = fitz.open(path)

        documents = []

        for page_num, page in enumerate(pdf):

            text = page.get_text()

            if text.strip():

                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": path.name,
                            "page": page_num + 1,
                            "file_type": ".pdf"
                        }
                    )
                )

        pdf.close()

        return documents

    def load_docx(self, file_path: str) -> List[Document]:
        """
        Load a DOCX file.
        """

        path = Path(file_path)

        doc = DocxDocument(path)

        text = "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
        )

        return [
            Document(
                page_content=text,
                metadata={
                    "source": path.name,
                    "file_type": ".docx"
                }
            )
        ]

    def load_document(self, file_path: str) -> List[Document]:
        """
        Automatically load a document
        based on its extension.
        """

        extension = Path(file_path).suffix.lower()

        if extension == ".txt":
            return self.load_txt(file_path)

        elif extension == ".pdf":
            return self.load_pdf(file_path)

        elif extension == ".docx":
            return self.load_docx(file_path)

        else:
            raise ValueError(
                f"Unsupported file type: {extension}"
            )